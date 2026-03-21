"""
resume_parser.py
~~~~~~~~~~~~~~~~
Robust resume text extraction and section parsing.

Key design decisions:
- PDF extraction works at the CHARACTER level using pdfplumber's char objects.
  This is the only reliable method for PDFs whose fonts omit space glyphs between
  words (common in design-tool exports like Canva, LaTeX, custom TeX fonts).
  Spaces are reconstructed from the x-coordinate gap between adjacent characters.
- A camelCase splitter with a protected-terms list handles any residual merged words.
- Section detection uses flexible regex that tolerates decorative characters around
  headings and recognises many synonym keywords per section.
"""

import re
from html.parser import HTMLParser
from typing import Dict, List, Optional
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

from fastapi import UploadFile
from PyPDF2 import PdfReader
from docx import Document

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None

try:
    from pdf2image import convert_from_bytes
    import pytesseract
except ImportError:  # pragma: no cover
    convert_from_bytes = None
    pytesseract = None


# ---------------------------------------------------------------------------
# Regex constants
# ---------------------------------------------------------------------------
EMAIL_REGEX = re.compile(r"[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}")
PHONE_REGEX = re.compile(r"\+?[\d\s\-\.\(\)]{7,}")
URL_REGEX = re.compile(r"(?:https?://)?(?:www\.)?[A-Za-z0-9\-_.]+\.[A-Za-z]{2,}(?:/[^\s<>]*)?")
MAILTO_REGEX = re.compile(r"mailto:([\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,})", re.IGNORECASE)
TEL_REGEX = re.compile(r"tel:([\d\s\-\.\(\)]+)", re.IGNORECASE)
LINKEDIN_SHORTHAND_RE = re.compile(r"linkedin\s*[:/]?\s*([\w\-]+(?:/[\w\-]+)*)", re.IGNORECASE)
GITHUB_SHORTHAND_RE = re.compile(r"github\s*[:/]?\s*([\w\-]+)", re.IGNORECASE)

# Terms that must never be split by the camelCase post-processor.
# Sorted longest-first so longer matches take priority.
_PROTECTED_TERMS: List[str] = sorted([
    "LLMOps", "MLOps", "DevOps", "FastAPI", "GitHub", "GitLab",
    "PyTorch", "TensorFlow", "OpenCV", "PostgreSQL", "LinkedIn",
    "LlamaIndex", "MLflow", "NumPy", "SciPy", "HuggingFace",
    "LangChain", "LangGraph", "AutoGen", "CrewAI", "scikit-learn",
    "pdfplumber", "DataFrame", "PoC", "MVP", "CI/CD",
    "APIs", "LLMs", "LLM", "RAG", "AI", "ML", "NLP",
    "CNN", "GAN", "GPU", "CPU", "AWS", "GCP", "DVC", "FAISS",
], key=len, reverse=True)


# ---------------------------------------------------------------------------
# URL / link utilities
# ---------------------------------------------------------------------------
def extract_urls(text: str) -> List[str]:
    return [u.strip().rstrip(".,;") for u in URL_REGEX.findall(text) if "@" not in u]


def _clean_url(url: str) -> str:
    url = url.strip().rstrip(".,;")
    if url.startswith("www."):
        return f"https://{url}"
    if re.match(r"^[A-Za-z0-9\-]{2,}\.[A-Za-z]{2,}($|/)", url):
        return f"https://{url}"
    return url


def extract_profile_links(text: str, extra_urls: Optional[List[str]] = None) -> Dict[str, str]:
    urls = extract_urls(text) + extract_urls(re.sub(r"\s+", "", text))
    if extra_urls:
        urls.extend(extra_urls)

    links: Dict[str, str] = {}
    for url in urls:
        url = _clean_url(url)
        low = url.lower()
        if "linkedin.com" in low and "in/" in low:
            links.setdefault("linkedin_url", url)
        elif "github.com" in low:
            links.setdefault("github_url", url)
        elif any(k in low for k in ("behance.net", "dribbble.com", "portfolio")):
            links.setdefault("portfolio_url", url)
        elif low.startswith("mailto:"):
            m = MAILTO_REGEX.search(url)
            if m:
                links.setdefault("email", m.group(1))
        elif low.startswith("tel:"):
            m = TEL_REGEX.search(url)
            if m:
                links.setdefault("phone", m.group(1))
        elif "http" in low and "linkedin" not in low and "github" not in low:
            links.setdefault("portfolio_url", url)

    if "linkedin_url" not in links:
        m = LINKEDIN_SHORTHAND_RE.search(text)
        if m:
            handle = m.group(1).strip().strip("/")
            if "@" not in handle and len(handle) > 2:
                links["linkedin_url"] = f"https://linkedin.com/in/{handle}"

    if "github_url" not in links:
        m = GITHUB_SHORTHAND_RE.search(text)
        if m:
            handle = m.group(1).strip()
            if "@" not in handle and len(handle) > 1:
                links["github_url"] = f"https://github.com/{handle}"

    return links


# ---------------------------------------------------------------------------
# HTML meta-tag parser (for portfolio URL metadata)
# ---------------------------------------------------------------------------
class _MetaParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.title: Optional[str] = None
        self.description: Optional[str] = None
        self._in_title = False

    def handle_starttag(self, tag, attrs):
        if tag.lower() == "title":
            self._in_title = True
        if tag.lower() == "meta":
            data = {k.lower(): v for k, v in attrs}
            if data.get("name") == "description" and data.get("content"):
                self.description = data["content"].strip()

    def handle_endtag(self, tag):
        if tag.lower() == "title":
            self._in_title = False

    def handle_data(self, data):
        if self._in_title:
            self.title = (self.title or "") + data.strip()


def fetch_url_metadata(url: str, timeout: int = 5) -> Dict[str, str]:
    url = url.strip()
    if url and not re.match(r"^https?://", url):
        url = f"https://{url}"
    try:
        req = Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urlopen(req, timeout=timeout) as resp:
            body = resp.read(100_000).decode(errors="ignore")
    except Exception:
        return {}
    parser = _MetaParser()
    parser.feed(body)
    return {k: v for k, v in {"title": parser.title, "description": parser.description}.items() if v}


# ---------------------------------------------------------------------------
# Text post-processing helpers
# ---------------------------------------------------------------------------
def _protect_terms(text: str) -> tuple:
    """Replace protected tech terms with null-byte placeholders."""
    placeholders: Dict[str, str] = {}
    for i, term in enumerate(_PROTECTED_TERMS):
        if term in text:
            ph = f"\x00{i:04d}\x00"
            placeholders[ph] = term
            text = text.replace(term, ph)
    return text, placeholders


def _restore_terms(text: str, placeholders: Dict[str, str]) -> str:
    for ph, term in placeholders.items():
        text = text.replace(ph, term)
    return text


def _insert_spaces_camel(text: str) -> str:
    """Insert spaces at lowercase→UpperCase word boundaries (camelCase artefacts).

    Protected tech terms (FastAPI, MLOps, etc.) are shielded from splitting.
    Only genuine English word boundaries are targeted.
    """
    text, phs = _protect_terms(text)
    # "buildingend" → "building end", "Proficientin" → "Proficient in"
    text = re.sub(r'([a-z]{2,})([A-Z][a-z])', r'\1 \2', text)
    # "deployment.Proficient" → "deployment. Proficient"
    text = re.sub(r'([,\.;:\)])([A-Za-z])', r'\1 \2', text)
    return _restore_terms(text, phs)


def _fix_hyphenation(text: str) -> str:
    """Join words split across lines by hyphenation: 'moni-\ntoring' → 'monitoring'."""
    return re.sub(r'(\w)-\n\s*(\w)', r'\1\2\n', text)


def _normalize_text(text: str) -> str:
    """Normalise text extracted from non-PDF sources (DOCX, TXT)."""
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\(cid:\d+\)", "", text)
    text = re.sub(r"[\u2600-\u26ff\u2700-\u27bf\uf000-\uffff]", " ", text)
    text = re.sub(r"[·•\u2022\u2023\u25E6\u2043\u2219]", "\n", text)
    lines = [re.sub(r" {2,}", " ", ln).strip() for ln in text.splitlines()]
    result = "\n".join(ln for ln in lines if ln)
    return re.sub(r"\n{3,}", "\n\n", result).strip()


def _match_label(text: str, patterns: List[str]) -> Optional[str]:
    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return None


# ---------------------------------------------------------------------------
# PDF extraction — character-level with gap-based space reconstruction
# ---------------------------------------------------------------------------
def _extract_text_from_pdf_chars(file: UploadFile) -> Optional[str]:
    """Extract PDF text at the character level, reconstructing spaces from x-gaps.

    Why char-level?
    Many PDF fonts — particularly those used by LaTeX/TeX and design tools like
    Canva — store no space glyph between words. The "space" is implied by the
    typesetter's positioning. Word-level and raw-text extraction merge such words
    into a single token. By iterating individual character objects and comparing
    the gap between consecutive x1 → x0, we can reliably insert spaces wherever
    the visual gap exceeds 20 % of the current font size.
    """
    if not pdfplumber:
        return None

    try:
        file.file.seek(0)
        all_lines: List[str] = []

        with pdfplumber.open(file.file) as pdf:
            for page in pdf.pages:
                chars = page.chars
                if not chars:
                    continue

                # Bucket characters by rounded y-position (same visual line)
                line_buckets: Dict[int, list] = {}
                for ch in chars:
                    y_key = round(ch["top"] / 2) * 2
                    line_buckets.setdefault(y_key, []).append(ch)

                for y_key in sorted(line_buckets):
                    line_chars = sorted(line_buckets[y_key], key=lambda c: c["x0"])
                    line = ""
                    prev_x1: Optional[float] = None
                    prev_size: Optional[float] = None

                    for ch in line_chars:
                        char = ch["text"]
                        x0, x1 = ch["x0"], ch["x1"]
                        size = float(ch.get("size") or 10)

                        # Actual space glyph — honour it directly
                        if char.strip() == "":
                            if line and not line.endswith(" "):
                                line += " "
                            prev_x1, prev_size = x1, size
                            continue

                        if prev_x1 is None:
                            line += char
                        else:
                            gap = x0 - prev_x1
                            # Gap > 20 % of font size = word boundary
                            if gap > (prev_size or size) * 0.20 and not line.endswith(" "):
                                line += " "
                            line += char

                        prev_x1, prev_size = x1, size

                    line = line.strip()
                    if line:
                        all_lines.append(line)

        text = "\n".join(all_lines)
        text = _fix_hyphenation(text)
        text = _insert_spaces_camel(text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()

        return text if len(text) > 100 else None

    except Exception:
        return None


def _ocr_pdf_text(file: UploadFile) -> str:
    """OCR fallback for scanned/image-only PDFs."""
    if not (convert_from_bytes and pytesseract):
        return ""
    file.file.seek(0)
    data = file.file.read()
    try:
        images = convert_from_bytes(data, dpi=200)
    except Exception:
        return ""
    texts: List[str] = []
    for image in images:
        try:
            texts.append(pytesseract.image_to_string(image, lang="eng"))
        except Exception:
            continue
    return _normalize_text("\n".join(texts))


def _extract_text_from_pdf(file: UploadFile) -> str:
    """Multi-strategy PDF text extraction.

    Priority:
    1. Char-level extraction (best for design-tool PDFs with missing space glyphs)
    2. pdfplumber word-level extraction (good fallback)
    3. PyPDF2 raw text (last text fallback)
    4. OCR (last resort for scanned PDFs)
    """
    # 1. Char-level — handles missing space glyphs
    result = _extract_text_from_pdf_chars(file)
    if result:
        return result

    # 2. pdfplumber word-level
    if pdfplumber:
        try:
            file.file.seek(0)
            with pdfplumber.open(file.file) as pdf:
                pages_text = []
                for page in pdf.pages:
                    words = page.extract_words(x_tolerance=2, y_tolerance=3)
                    line_buckets: Dict[int, list] = {}
                    for w in words:
                        y = round(w["top"] / 3) * 3
                        line_buckets.setdefault(y, []).append(w)
                    lines = []
                    for y in sorted(line_buckets):
                        ws = sorted(line_buckets[y], key=lambda w: w["x0"])
                        prev_x1 = None
                        line = ""
                        for w in ws:
                            if prev_x1 and w["x0"] - prev_x1 > 3:
                                line += " "
                            line += w["text"]
                            prev_x1 = w["x1"]
                        if line.strip():
                            lines.append(line.strip())
                    pages_text.append("\n".join(lines))
            text = _insert_spaces_camel(_fix_hyphenation("\n".join(pages_text)))
            if len(text) > 100:
                return text
        except Exception:
            pass

    # 3. PyPDF2 raw text
    try:
        file.file.seek(0)
        reader = PdfReader(file.file)
        raw = "\n".join(p.extract_text() or "" for p in reader.pages)
        text = _insert_spaces_camel(_fix_hyphenation(_normalize_text(raw)))
        if len(text) > 50:
            return text
    except Exception:
        pass

    # 4. OCR
    return _ocr_pdf_text(file)


# ---------------------------------------------------------------------------
# PDF / DOCX link extraction
# ---------------------------------------------------------------------------
def extract_pdf_links(file: UploadFile) -> List[str]:
    file.file.seek(0)
    urls: List[str] = []
    try:
        reader = PdfReader(file.file)
        for page in reader.pages:
            annots = page.get("/Annots")
            if not annots:
                continue
            for annot in annots:
                obj = annot.get_object()
                if not isinstance(obj, dict):
                    continue
                action = obj.get("/A")
                if action and isinstance(action, dict):
                    uri = action.get("/URI")
                    if uri:
                        urls.append(str(uri))
                elif obj.get("/URI"):
                    urls.append(str(obj["/URI"]))
    except Exception:
        pass
    return urls


def extract_docx_links(file: UploadFile) -> List[str]:
    file.file.seek(0)
    try:
        doc = Document(file.file)
        return [rel.target_ref for rel in doc.part.rels.values() if "hyperlink" in rel.reltype]
    except Exception:
        return []


# ---------------------------------------------------------------------------
# Public file-level APIs
# ---------------------------------------------------------------------------
def extract_text_from_resume(file: UploadFile) -> str:
    filename = (file.filename or "").lower()
    if filename.endswith(".pdf"):
        return _extract_text_from_pdf(file)
    if filename.endswith(".docx"):
        file.file.seek(0)
        doc = Document(file.file)
        return _normalize_text("\n".join(p.text for p in doc.paragraphs))
    # TXT or unknown
    file.file.seek(0)
    raw = file.file.read()
    if isinstance(raw, bytes):
        raw = raw.decode(errors="ignore")
    return _normalize_text(raw)


def extract_links_from_resume(file: UploadFile) -> List[str]:
    filename = (file.filename or "").lower()
    if filename.endswith(".pdf"):
        return extract_pdf_links(file)
    if filename.endswith(".docx"):
        return extract_docx_links(file)
    return []


# ---------------------------------------------------------------------------
# Field extractors
# ---------------------------------------------------------------------------
def extract_email(text: str) -> Optional[str]:
    m = MAILTO_REGEX.search(text)
    if m:
        return m.group(1)
    labeled = _match_label(text, [r"e-?mail\s*[:\-]\s*(\S+@\S+)"])
    if labeled:
        return labeled
    for src in (re.sub(r"\s+", "", text), text):
        m = EMAIL_REGEX.search(src)
        if m:
            return m.group(0)
    return None


def extract_phone(text: str) -> Optional[str]:
    m = TEL_REGEX.search(text)
    if m:
        return m.group(1).strip()
    labeled = _match_label(text, [r"(?:phone|tel|mobile|cell)\s*[:\-]\s*([+\d][\d\s\-\.\(\)]+)"])
    if labeled:
        return labeled
    for src in (re.sub(r"\s+", "", text), text):
        m = PHONE_REGEX.search(src)
        if m:
            digits = re.sub(r"\D", "", m.group(0))
            if len(digits) >= 7:
                return re.sub(r"\s+", " ", m.group(0)).strip()
    return None


def _trunc(value: str, max_len: int) -> str:
    return value.strip()[:max_len] if value else value


def extract_name(text: str) -> Optional[str]:
    labeled = _match_label(text, [r"(?:name|full name)\s*[:\-]\s*(.+)"])
    if labeled:
        return _trunc(labeled, 255)

    skip = {"email", "phone", "tel", "linkedin", "github", "resume", "cv",
             "profile", "summary", "experience", "education", "skills"}
    for line in [ln.strip() for ln in text.splitlines() if ln.strip()][:12]:
        if any(s in line.lower() for s in skip):
            continue
        words = line.split()
        if 2 <= len(words) <= 5 and sum(1 for w in words if w and w[0].isupper()) >= len(words) - 1:
            return _trunc(line, 255)
    return None


def extract_location(text: str) -> Optional[str]:
    labeled = _match_label(text, [
        r"location\s*[:\-]\s*(.+)",
        r"address\s*[:\-]\s*(.+)",
        r"based in\s+([^\n]+)",
    ])
    if labeled:
        return _trunc(labeled, 255)

    skip = {"email", "phone", "linkedin", "github", "www.", "experience",
            "education", "skills", "university", "http"}
    for line in text.splitlines():
        line = line.strip()
        if not line or len(line) > 80 or any(s in line.lower() for s in skip):
            continue
        if re.search(r"[A-Za-z\s\-]{2,},\s*[A-Za-z\s]{2,}", line):
            return _trunc(line, 255)
    return None


# ---------------------------------------------------------------------------
# Section parser
# ---------------------------------------------------------------------------
SECTION_KEYWORDS: Dict[str, List[str]] = {
    "summary": [
        "summary", "professional summary", "profile", "about me",
        "objective", "career objective", "personal statement", "about",
    ],
    "skills": [
        "skills", "technical skills", "skillset", "core competencies",
        "technologies", "tech stack", "tools", "key skills",
        "areas of expertise", "expertise",
    ],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment history", "employment", "work history",
        "career history", "professional background", "positions held",
        "internships", "internship experience",
    ],
    "education": [
        "education", "academic background", "qualifications",
        "academic qualifications", "educational background",
        "degrees", "academic history",
    ],
    "projects": [
        "projects", "selected projects", "key projects", "personal projects",
        "academic projects", "side projects", "notable projects",
        "project experience",
    ],
    "certifications": [
        "certifications", "certificates", "licenses", "accreditations",
        "professional certifications", "courses", "training",
        "professional development", "certifications & additional",
    ],
    "languages": [
        "languages", "language skills", "spoken languages",
        "language proficiency", "linguistic skills",
    ],
}


def _build_heading_re() -> re.Pattern:
    phrases = sorted(
        (p for ps in SECTION_KEYWORDS.values() for p in ps),
        key=len, reverse=True,
    )
    alt = "|".join(re.escape(p) for p in phrases)
    return re.compile(
        r"^[\s\W]*(?P<kw>" + alt + r")[\s\W]*$",
        flags=re.IGNORECASE | re.MULTILINE,
    )


def _resolve_key(keyword: str) -> Optional[str]:
    kw = keyword.strip().lower()
    for key, phrases in SECTION_KEYWORDS.items():
        if kw in phrases:
            return key
    return None


def parse_resume_sections(text: str) -> Dict[str, str]:
    """Parse resume sections robustly.

    - Headings with decorative characters (──, ==, ···) are matched
    - Many synonym keywords per section
    - Duplicate section blocks are merged (e.g. two Experience headings)
    - Bullet/dash prefixes are stripped from content lines
    """
    heading_re = _build_heading_re()
    matches = [m for m in heading_re.finditer(text) if len(m.group("kw").strip()) >= 3]
    if not matches:
        return {}

    sections: Dict[str, str] = {}
    for idx, match in enumerate(matches):
        key = _resolve_key(match.group("kw"))
        if not key:
            continue

        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        raw = text[start:end].strip()

        lines = []
        for line in raw.splitlines():
            line = re.sub(r"^[\s\-\*\•\>\~\=\|·]+", "", line).strip()
            if line:
                lines.append(line)

        content = "\n".join(lines).strip()
        if not content:
            continue

        # Merge if section key already seen (handles duplicate headings)
        sections[key] = (sections[key] + "\n\n" + content) if key in sections else content

    return sections